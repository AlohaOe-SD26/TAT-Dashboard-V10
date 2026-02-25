# src/api/mis_automation.py — v2.0
# ─────────────────────────────────────────────────────────────────────────────
# MIS Automation routes: browser init, deal creation, end-date updates,
# validation injection, and pre-flight validation.
# Selenium ops live in src/automation/mis_entry.py (no-touch zone).
# ─────────────────────────────────────────────────────────────────────────────

from __future__ import annotations

import os
import sys
import traceback

from flask import Blueprint, jsonify, request

from src.session import session

bp = Blueprint('mis_automation', __name__)


# ── System ────────────────────────────────────────────────────────────────────

@bp.route('/api/restart', methods=['POST'])
def restart():
    """Hard restart via os.execv."""
    try:
        print("\n" + "="*70 + "\n[RESTART] Restarting application...\n" + "="*70 + "\n")
        os.execv(sys.executable, [sys.executable] + sys.argv)
        return jsonify({'success': True})
    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})


@bp.route('/api/browser-status')
def browser_status():
    """Return current browser readiness."""
    return jsonify({
        'ready':    session.is_browser_ready(),
        'instance': session.get_browser() is not None,
    })


@bp.route('/api/get-settings-dropdowns', methods=['GET'])
def get_settings_dropdowns():
    """Fetch dropdown options from Settings tab for Enhanced Create Popup."""
    try:
        from src.integrations.google_sheets import load_settings_dropdown_data

        spreadsheet_id = session.get_spreadsheet_id()
        if not spreadsheet_id:
            return jsonify({'success': False, 'error': 'No spreadsheet loaded. Select a Google Sheet first.'})

        service = session.get_sheets_service()
        if not service:
            return jsonify({'success': False, 'error': 'Sheets service not authenticated.'})

        data = load_settings_dropdown_data(spreadsheet_id, service)
        return jsonify({
            'success':         True,
            'stores':          data.get('stores', []),
            'categories':      data.get('categories', []),
            'brand_linked_map': data.get('brand_linked_map', {}),
        })
    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})


@bp.route('/api/init-all', methods=['POST'])
def init_all():
    """Initialize browser, MIS login, and Blaze login in sequence."""
    try:
        from src.automation.browser import init_browser, mis_login, robust_login
        from src.integrations.blaze_api import load_stored_token, validate_token

        data       = request.get_json() or {}
        mis_creds  = data.get('mis', {})
        blaze_creds = data.get('blaze', {})

        session.set_mis_credentials(mis_creds)
        session.set_blaze_credentials(blaze_creds)

        driver = session.get_browser()

        # 1. Browser — guard: need both the flag AND a live driver object in memory.
        # browser_ready can be stale-True in SQLite from a previous run while the
        # actual WebDriver is gone, so we check both.
        driver_alive = False
        if driver is not None:
            try:
                _ = driver.window_handles  # lightweight liveness check
                driver_alive = True
            except Exception:
                # Driver is dead (window was closed, Chrome crashed, etc.)
                session.set_browser(None)
                driver = None

        if not driver_alive:
            driver = init_browser()
            if not driver:
                return jsonify({'success': False, 'error': 'Failed to initialize browser'})
            session.set_browser(driver)
            session.set_browser_ready(True)

        # 2. MIS Login
        mis_success = False
        if mis_creds.get('username') and mis_creds.get('password'):
            mis_success = mis_login(driver, mis_creds['username'], mis_creds['password'], new_tab=True)

        # 3. Blaze Login
        blaze_success = False
        if blaze_creds.get('email') and blaze_creds.get('password'):
            stored_token = load_stored_token()
            if validate_token(stored_token):
                print("[INIT] Existing token valid. Skipping sniffer.")
                robust_login(blaze_creds['email'], blaze_creds['password'])
                session.set_blaze_token(stored_token)
                blaze_success = True
            else:
                login_status = robust_login(blaze_creds['email'], blaze_creds['password'])
                blaze_success = bool(login_status)

        return jsonify({
            'success':        True,
            'browser_ready':  session.is_browser_ready(),
            'mis_login':      mis_success,
            'blaze_login':    blaze_success,
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})


# ── Deal Automation ───────────────────────────────────────────────────────────

@bp.route('/api/mis/create-deal', methods=['POST'])
def create_deal():
    """Automation: Fill MIS modal via Selenium. Builds ValidationRecord for pre-flight."""
    try:
        from src.automation.mis_entry import fill_deal_form
        from src.core.validation_engine import engine as ve, ValidationRecord

        data   = request.get_json() or {}
        driver = session.get_browser()

        if not driver:
            return jsonify({'success': False, 'error': 'Browser not initialized'})

        result = fill_deal_form(driver, data)
        if not result.get('success'):
            return jsonify(result)

        # Pre-flight: validate entered vs expected data
        if data.get('expected'):
            source = ValidationRecord(**{
                k: str(result.get('filled', {}).get(k, ''))
                for k in ValidationRecord.__dataclass_fields__
            })
            target = ValidationRecord(**{
                k: str(data['expected'].get(k, ''))
                for k in ValidationRecord.__dataclass_fields__
            })
            field_results = ve.compare(source, target)
            summary       = ve.summary(field_results)
            result['validation'] = summary.to_dict()

        return jsonify(result)

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})


@bp.route('/api/mis/automate-create-deal', methods=['POST'])
def automate_create_deal():
    """Full automation path for Up-Down Planning and ID Matcher create buttons."""
    try:
        from src.automation.mis_entry import automate_full_create

        data   = request.get_json() or {}
        driver = session.get_browser()

        if not driver:
            return jsonify({'success': False, 'error': 'Browser not initialized'})

        result = automate_full_create(driver, data, session=session)
        return jsonify(result)

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})


@bp.route('/api/mis/update-end-date', methods=['POST'])
def update_end_date():
    """Automation: Update an existing MIS entry's end date via Selenium."""
    try:
        from src.automation.mis_entry import update_mis_end_date

        data   = request.get_json() or {}
        driver = session.get_browser()

        if not driver:
            return jsonify({'success': False, 'error': 'Browser not initialized'})

        result = update_mis_end_date(driver, data)
        return jsonify(result)

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})


@bp.route('/api/mis/automate-end-date', methods=['POST'])
def automate_end_date():
    """Full end-date automation sequence."""
    try:
        from src.automation.mis_entry import automate_full_end_date

        data   = request.get_json() or {}
        driver = session.get_browser()

        if not driver:
            return jsonify({'success': False, 'error': 'Browser not initialized'})

        result = automate_full_end_date(driver, data, session=session)
        return jsonify(result)

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})


@bp.route('/api/mis/inject-validation', methods=['POST'])
def inject_validation():
    """Inject MIS validation system into the current browser page."""
    try:
        from src.automation.browser import inject_mis_validation

        driver = session.get_browser()
        if not driver:
            return jsonify({'success': False, 'error': 'Browser not initialized'})

        # Switch to MIS tab
        for handle in driver.window_handles:
            driver.switch_to.window(handle)
            if 'mymis.net' in driver.current_url or 'mis' in driver.current_url.lower():
                break

        inject_mis_validation(driver, expected_data=None)
        return jsonify({'success': True, 'message': 'Validation system injected. Modal monitoring active.'})

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})


@bp.route('/api/mis/open-sheet-row', methods=['POST'])
def open_sheet_row():
    """Open a specific Google Sheet row in the browser."""
    try:
        from src.automation.browser import open_google_sheet_in_browser

        data           = request.get_json() or {}
        spreadsheet_id = session.get_spreadsheet_id()
        sheet_name     = data.get('sheet_name') or session.get_mis_current_sheet()
        row_num        = data.get('row')

        if not spreadsheet_id or not sheet_name:
            return jsonify({'success': False, 'error': 'No spreadsheet/sheet configured'})

        driver = session.get_browser()
        if not driver:
            return jsonify({'success': False, 'error': 'Browser not initialized'})

        success = open_google_sheet_in_browser(spreadsheet_id, sheet_name, row_num)
        return jsonify({'success': success})

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)})


# ── Utilities ─────────────────────────────────────────────────────────────────

@bp.route('/api/mis/generate-newsletter', methods=['POST'])
def generate_newsletter():
    """
    Generate Newsletter files:
    1. Excel file with 6 tabs (CLUB420 and TAT LEGACY × Weekly/Monthly/Sale)
    2. Two DOCX files (CLUB420_Newsletter.docx, TAT_LEGACY_Newsletter.docx)
    Each DOCX has 3 tables with day-of-week columns.
    Monolith: lines 24528–24938.
    """
    import re as _re
    import traceback as _tb
    import pandas as _pd
    from pathlib import Path as _Path
    from datetime import datetime as _dt

    try:
        # ── Session source ────────────────────────────────────────────────────
        sections = session.get('mis_generated_sections') or {}
        if not sections:
            return jsonify({'success': False, 'error': "No CSV generated. Please click 'Generate CSV' first."}), 400

        _PROJECT_ROOT   = _Path(__file__).resolve().parent.parent.parent
        _NEWSLETTER_DIR = _PROJECT_ROOT / 'reports' / 'NEWSLETTERS'
        _NEWSLETTER_DIR.mkdir(parents=True, exist_ok=True)

        timestamp    = _dt.now().strftime('%Y%m%d_%H%M%S')
        current_month = _dt.now().strftime('%B')

        newsletter_folder = _NEWSLETTER_DIR / f'NEWSLETTER_{timestamp}'
        newsletter_folder.mkdir(parents=True, exist_ok=True)
        print(f"[NEWSLETTER] Output folder: {newsletter_folder}")

        # ── Location helpers ──────────────────────────────────────────────────
        CLUB420_LOCS = ['Davis', 'Dixon']
        DD_ONLY_PATTERNS = {
            'Davis', 'Dixon', 'Davis, Dixon', 'Dixon, Davis',
            'All Locations Except: Davis, Dixon',
            'All Locations Except: Dixon, Davis',
            'All Locations Except: Davis',
            'All Locations Except: Dixon',
        }

        def _is_club420(store: str) -> bool:
            if not store: return False
            s = str(store).strip()
            if s == 'All Locations': return True
            is_except = 'all locations except' in s.lower() or 'except' in s.lower()
            for loc in CLUB420_LOCS:
                if loc in s:
                    return not is_except
            return False

        def _is_tat_legacy(store: str) -> bool:
            if not store: return False
            s = str(store).strip()
            if s == 'All Locations': return True
            if s in DD_ONLY_PATTERNS: return False
            if all(p in CLUB420_LOCS for p in [x.strip() for x in s.split(',')]): return False
            if 'all locations except' in s.lower(): return True
            return True

        # ── Row formatters ────────────────────────────────────────────────────
        def _fmt_discount(row: dict) -> str:
            ui = row.get('UI_DEAL_INFO', [])
            info = ''
            if isinstance(ui, list) and ui:
                first = ui[0]
                info  = str(first.get('info', '') if isinstance(first, dict) else first).strip()
            if not info:
                info = str(row.get('Deal Information', '') or row.get('Deal Info', '')).strip()
            iup = info.upper()
            if 'BOGO FOR $1' in iup or 'B2G1 FOR $1' in iup:
                m = _re.search(r'(BOGO\s+for\s+\$\d+|B2G1\s+for\s+\$\d+)', info, _re.IGNORECASE)
                return m.group(0) if m else ('BOGO for $1' if 'BOGO' in iup else 'B2G1 for $1')
            disc = row.get('Daily Deal Discount', '')
            if disc:
                ds = str(disc).strip().replace('%', '').strip()
                try:
                    n = float(ds)
                    return f"{int(n)}% Off" if n == int(n) else f"{n}% Off"
                except (ValueError, TypeError):
                    return f"{ds}% Off" if ds else ''
            return str(disc) if disc else ''

        def _fmt_deal_info(discount: str, cats: str) -> str:
            if not discount: return ''
            d_up = discount.upper()
            if 'BOGO' in d_up or 'B2G1' in d_up:
                return discount if 'Like-Products' in discount else f"{discount} - Like-Products"
            if not cats or cats.strip().lower() == 'all categories':
                return f"{discount} all Products"
            return f"{discount} {cats}"

        def _parse_days_weekly(wday: str) -> dict:
            days = {k: '' for k in ('Mon.', 'Tues.', 'Wed.', 'Thurs.', 'Fri.', 'Sat.', 'Sun.')}
            MAP  = {
                'monday': ('Mon.', 'Mon'), 'tuesday': ('Tues.', 'Tues'),
                'wednesday': ('Wed.', 'Wed'), 'thursday': ('Thurs.', 'Thurs'),
                'friday': ('Fri.', 'Fri'), 'saturday': ('Sat.', 'Sat'),
                'sunday': ('Sun.', 'Sun'),
            }
            wl = (wday or '').lower()
            for full, (key, abbr) in MAP.items():
                if full in wl:
                    days[key] = abbr
            return days

        def _parse_days_monthly(wday: str) -> dict:
            days = {k: '' for k in ('Mon.', 'Tues.', 'Wed.', 'Thurs.', 'Fri.', 'Sat.', 'Sun.')}
            if wday:
                days['Mon.'] = str(wday).strip()
            return days

        def _process_row(row: dict, sec: str) -> dict:
            store  = row.get('Store', '') or row.get('DISPLAY_STORE', '') or ''
            disc   = _fmt_discount(row)
            cats   = row.get('Category', '') or row.get('DISPLAY_CATEGORY', '')
            wday   = row.get('Weekday', '')
            brand  = row.get('Brand', '')
            days   = _parse_days_monthly(wday) if sec == 'monthly' else _parse_days_weekly(wday)
            return {
                'Weekday': wday, 'Brand': brand, 'Discount': disc,
                'Categories': cats, 'Deal_Info': _fmt_deal_info(disc, cats),
                'Days': days, 'Store': store,
            }

        # ── Build section_data ────────────────────────────────────────────────
        section_data: dict = {
            'weekly':  {'club420': [], 'tat_legacy': []},
            'monthly': {'club420': [], 'tat_legacy': []},
            'sale':    {'club420': [], 'tat_legacy': []},
        }
        for skey in ('weekly', 'monthly', 'sale'):
            for raw_row in sections.get(skey, {}).get('rows', []):
                p = _process_row(raw_row, skey)
                s = p['Store']
                if _is_club420(s):    section_data[skey]['club420'].append(p)
                if _is_tat_legacy(s): section_data[skey]['tat_legacy'].append(p)

        # ── Excel (6 tabs) ────────────────────────────────────────────────────
        excel_path = newsletter_folder / f'Newsletter_Table_{timestamp}.xlsx'
        col_order  = ['Weekday', 'Brand', 'Discount', 'Categories']
        with _pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
            for skey, title in (('weekly', 'Weekly Deals'), ('monthly', 'Monthly Deals'), ('sale', 'Sale Deals')):
                for loc_key, loc_label in (('club420', 'CLUB420'), ('tat_legacy', 'TAT LEGACY')):
                    rows = section_data[skey][loc_key]
                    df   = _pd.DataFrame([{k: r[k] for k in col_order} for r in rows]) if rows else _pd.DataFrame(columns=col_order)
                    for c in col_order:
                        if c not in df.columns: df[c] = ''
                    df[col_order].to_excel(writer, sheet_name=f'{loc_label} - {title}', index=False)
        print(f"[NEWSLETTER] Excel generated: {excel_path.name}")

        # ── DOCX generation ───────────────────────────────────────────────────
        try:
            from docx import Document as _DocxDoc
            from docx.shared import Inches as _Inches, Pt as _Pt, RGBColor as _RGB
            from docx.enum.text import WD_ALIGN_PARAGRAPH as _WALIGN
            from docx.enum.table import WD_TABLE_ALIGNMENT as _TALIGN
            from docx.oxml.ns import nsdecls as _nsdecls
            from docx.oxml import parse_xml as _parse_xml
            _DOCX_AVAILABLE = True
        except ImportError:
            print("[NEWSLETTER] python-docx not installed — skipping DOCX generation")
            _DOCX_AVAILABLE = False

        def _create_docx(loc_name: str, loc_data: dict, out_path: _Path) -> None:
            if not _DOCX_AVAILABLE:
                return
            doc = _DocxDoc()
            for sec in doc.sections:
                for attr in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
                    setattr(sec, attr, _Inches(0.5))
            t = doc.add_paragraph()
            t.alignment = _WALIGN.CENTER
            tr = t.add_run(loc_name); tr.bold = True; tr.font.size = _Pt(18); tr.font.name = 'Arial'
            s = doc.add_paragraph()
            s.alignment = _WALIGN.CENTER
            sr = s.add_run(f"{current_month} Daily Deal Schedule Notes:")
            sr.bold = True; sr.font.size = _Pt(14); sr.font.name = 'Arial'
            doc.add_paragraph()

            def _add_table(sec_title: str, rows: list) -> None:
                p = doc.add_paragraph()
                r = p.add_run(sec_title); r.bold = True; r.font.size = _Pt(12); r.font.name = 'Arial'
                if not rows:
                    np2 = doc.add_paragraph()
                    nr  = np2.add_run("No deals in this section.")
                    nr.italic = True; nr.font.size = _Pt(10); nr.font.name = 'Arial'
                    nr.font.color.rgb = _RGB(102, 102, 102)
                    doc.add_paragraph(); return
                tbl = doc.add_table(rows=1 + len(rows), cols=9)
                tbl.style = 'Table Grid'; tbl.alignment = _TALIGN.CENTER
                widths = [_Inches(1.2), _Inches(2.3)] + [_Inches(0.5)] * 7
                for row_obj in tbl.rows:
                    for ci, cell in enumerate(row_obj.cells):
                        cell.width = widths[ci]
                hdrs = ['Brand', 'Deal Info', 'Mon.', 'Tues.', 'Wed.', 'Thurs.', 'Fri.', 'Sat.', 'Sun.']
                for ci, hdr in enumerate(hdrs):
                    cell = tbl.rows[0].cells[ci]; cell.text = ''
                    para = cell.paragraphs[0]; para.alignment = _WALIGN.CENTER
                    run  = para.add_run(hdr); run.bold = True; run.font.size = _Pt(9); run.font.name = 'Arial'
                    cell._tc.get_or_add_tcPr().append(_parse_xml(f'<w:shd {_nsdecls("w")} w:fill="D9E2F3"/>'))
                for ri, rd in enumerate(rows):
                    drow = tbl.rows[ri + 1]; days = rd.get('Days', {})
                    vals = [rd.get('Brand', ''), rd.get('Deal_Info', '')] + [days.get(h, '') for h in hdrs[2:]]
                    for ci, val in enumerate(vals):
                        cell = drow.cells[ci]; cell.text = ''
                        para = cell.paragraphs[0]
                        if ci >= 2: para.alignment = _WALIGN.CENTER
                        run = para.add_run(str(val)); run.font.size = _Pt(9); run.font.name = 'Arial'
                doc.add_paragraph()

            _add_table("Weekly Deals",  loc_data.get('weekly',  []))
            _add_table("Monthly Deals", loc_data.get('monthly', []))
            _add_table("Sale Deals",    loc_data.get('sale',    []))
            doc.save(str(out_path))
            print(f"[NEWSLETTER] DOCX saved: {out_path.name}")

        club420_docx    = newsletter_folder / f'CLUB420_Newsletter_{timestamp}.docx'
        tat_legacy_docx = newsletter_folder / f'TAT_LEGACY_Newsletter_{timestamp}.docx'
        try:
            _create_docx("CLUB420",     {'weekly': section_data['weekly']['club420'],
                                          'monthly': section_data['monthly']['club420'],
                                          'sale':    section_data['sale']['club420']}, club420_docx)
        except Exception as e:
            print(f"[NEWSLETTER] CLUB420 DOCX error: {e}"); _tb.print_exc()
        try:
            _create_docx("TAT LEGACY",  {'weekly': section_data['weekly']['tat_legacy'],
                                          'monthly': section_data['monthly']['tat_legacy'],
                                          'sale':    section_data['sale']['tat_legacy']}, tat_legacy_docx)
        except Exception as e:
            print(f"[NEWSLETTER] TAT LEGACY DOCX error: {e}"); _tb.print_exc()

        # ── Summary log ───────────────────────────────────────────────────────
        print(f"[NEWSLETTER] ===== GENERATION COMPLETE =====")
        print(f"[NEWSLETTER] Excel:      {excel_path.exists()}")
        print(f"[NEWSLETTER] CLUB420:    {club420_docx.exists()}")
        print(f"[NEWSLETTER] TAT LEGACY: {tat_legacy_docx.exists()}")

        return jsonify({
            'success': True,
            'folder':  str(newsletter_folder),
            'files': {
                'excel':          str(excel_path)         if excel_path.exists()    else None,
                'club420_docx':   str(club420_docx)       if club420_docx.exists()  else None,
                'tat_legacy_docx': str(tat_legacy_docx)   if tat_legacy_docx.exists() else None,
            },
            'counts': {
                loc: {sec: len(section_data[sec][loc]) for sec in ('weekly', 'monthly', 'sale')}
                for loc in ('club420', 'tat_legacy')
            },
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500


# NOTE: /api/tax-rates and /api/save-tax-rates are owned by src/api/blaze.py.
# Removed from here to prevent duplicate URL rule conflicts.
